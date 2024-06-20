import os
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

def extract_epikriser(doc_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if not os.path.isfile(doc_path):
        logging.error(f"The file '{doc_path}' does not exist.")
        return

    doc = Document(doc_path)
    current_title = None
    current_content = []
    title_counter = {}

    def process_paragraphs(paragraphs):
        nonlocal current_title, current_content
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                logging.debug("Encountered an empty paragraph, skipping.")
                continue

            logging.debug(f"Processing paragraph: {text}")

            if text.startswith('T:'):
                if current_title and current_content:
                    save_to_markdown(current_title, current_content, output_dir)
                    logging.info(f"Saved document: {current_title}")
                current_title = text[2:].strip()
                if current_title in title_counter:
                    title_counter[current_title] += 1
                    current_title = f"{current_title} {title_counter[current_title]}"
                else:
                    title_counter[current_title] = 1
                current_content = []
                logging.info(f"Started new document: {current_title}")
            elif text.startswith('S:'):
                section_title = text[2:].strip()
                content_split = section_title.split('C:', 1)
                if len(content_split) > 1:
                    section_title, section_content = content_split
                    current_content.append(f'## {section_title.strip()}\n')
                    current_content.append(f'{section_content.strip()}\n\n')
                else:
                    current_content.append(f'## {section_title.strip()}\n')
            elif text.startswith('C:'):
                current_content.append(text[2:].strip() + '\n\n')
            elif text.startswith('P:'):
                current_content.append(f'- {text[2:].strip()}\n')

    for para in doc.paragraphs:
        process_paragraphs([para])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    if current_title and current_content:
        save_to_markdown(current_title, current_content, output_dir)
        logging.info(f"Saved document: {current_title}")

def save_to_markdown(title, content, output_dir):
    file_name = f"{title.replace(' ', '_')}.md"
    file_path = os.path.join(output_dir, file_name)
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(f"# {title}\n")
        f.writelines(content)
    logging.info(f"Generated file: {file_path}")

if __name__ == "__main__":
    # Use a relative path to the script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, 'epikriser.docx')
    output_dir = os.path.join(script_dir, 'output_md_files')

    logging.info(f"Document path: {doc_path}")
    logging.info(f"Output directory: {output_dir}")

    extract_epikriser(doc_path, output_dir)
    logging.info(f"Markdown files have been saved to the '{output_dir}' directory.")
